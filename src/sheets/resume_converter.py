import os
import tempfile
import requests
from typing import Optional
from pathlib import Path

class ResumeConverter:
    """Converts various resume formats to PDF using proper libraries"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.jpg', '.jpeg', '.png']
    
    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """
        Extract text directly from resume file
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted text content, or None if extraction fails
        """
        try:
            file_ext = Path(file_path).suffix.lower()
            
            # First, try to detect the actual file type by examining content
            actual_file_type = self._detect_file_type(file_path)
            
            if actual_file_type == 'pdf':
                return self._extract_text_from_pdf(file_path)
            elif actual_file_type == 'docx':
                return self._extract_text_from_docx(file_path)
            elif actual_file_type == 'doc':
                return self._extract_text_from_doc(file_path)
            elif actual_file_type == 'txt':
                return self._extract_text_from_txt(file_path)
            elif actual_file_type == 'html':
                return self._extract_text_from_html(file_path)
            else:
                # Fallback to extension-based detection
                if file_ext == '.pdf':
                    return self._extract_text_from_pdf(file_path)
                elif file_ext == '.docx':
                    return self._extract_text_from_docx(file_path)
                elif file_ext == '.doc':
                    return self._extract_text_from_doc(file_path)
                elif file_ext == '.txt':
                    return self._extract_text_from_txt(file_path)
                else:
                    print(f"Unsupported file format: {file_ext}")
                    return None
                
        except Exception as e:
            print(f"Error extracting text from file: {e}")
            return None
    
    def _extract_text_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text directly from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            result = '\n'.join(text_content)
            print(f"Successfully extracted {len(result)} characters from DOCX")
            return result
            
        except ImportError:
            print("python-docx not available")
            return None
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return None
    
    def _extract_text_from_doc(self, file_path: str) -> Optional[str]:
        """Extract text directly from DOC file"""
        try:
            # Try docx2txt first (works for many DOC files)
            try:
                import docx2txt
                text_content = docx2txt.process(file_path)
                print(f"Successfully extracted {len(text_content)} characters from DOC using docx2txt")
                return text_content
            except Exception as e:
                print(f"docx2txt failed: {e}")
            
            # If that fails, try basic text extraction
            try:
                with open(file_path, 'rb') as f:
                    content = f.read()
                
                # Simple text extraction
                text_content = ""
                for byte in content:
                    if 32 <= byte <= 126 or byte in [9, 10, 13]:  # Printable ASCII + whitespace
                        text_content += chr(byte)
                
                # Clean up the text
                lines = text_content.split('\n')
                cleaned_lines = []
                for line in lines:
                    if len(line.strip()) > 0 and not line.strip().startswith('\x00'):
                        cleaned_line = ''.join(char for char in line if ord(char) >= 32 or char in '\n\t')
                        if cleaned_line.strip():
                            cleaned_lines.append(cleaned_line.strip())
                
                text_content = '\n'.join(cleaned_lines[:50])  # Limit to first 50 lines
                print(f"Successfully extracted {len(text_content)} characters from DOC using basic method")
                return text_content
                
            except Exception as e:
                print(f"Basic text extraction failed: {e}")
                return None
                
        except Exception as e:
            print(f"Error extracting text from DOC: {e}")
            return None
    
    def _extract_text_from_txt(self, file_path: str) -> Optional[str]:
        """Extract text directly from TXT file"""
        try:
            # Read text content with multiple encoding attempts
            content = None
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        content = f.read()
                    print(f"Successfully read TXT file with {encoding} encoding")
                    break
                except Exception as e:
                    print(f"Failed to read with {encoding}: {e}")
                    continue
            
            if not content:
                print("Could not read TXT file with any encoding")
                return None
            
            # Clean up the content - remove excessive whitespace
            lines = content.split('\n')
            cleaned_lines = []
            for line in lines:
                if line.strip():
                    cleaned_lines.append(line.strip())
            
            # Limit to first 100 lines to avoid massive files
            cleaned_lines = cleaned_lines[:100]
            
            if not cleaned_lines:
                print("No meaningful content found in TXT file")
                return None
            
            result = '\n'.join(cleaned_lines)
            print(f"Successfully extracted {len(result)} characters from TXT")
            return result
            
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return None
    
    def _extract_text_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text directly from PDF file"""
        try:
            import PyPDF2
            
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(page_text.strip())
                            print(f"Page {page_num + 1}: Extracted {len(page_text)} characters")
                        else:
                            print(f"Page {page_num + 1}: No text extracted - trying OCR...")
                            # Try OCR for image-based PDFs
                            ocr_text = self._extract_text_with_ocr(file_path, page_num)
                            if ocr_text:
                                text_content.append(ocr_text)
                                print(f"Page {page_num + 1}: OCR extracted {len(ocr_text)} characters")
                            else:
                                print(f"Page {page_num + 1}: OCR also failed")
                    except Exception as e:
                        print(f"Error extracting text from page {page_num + 1}: {e}")
                        # Try OCR as fallback
                        print(f"Page {page_num + 1}: Trying OCR as fallback...")
                        ocr_text = self._extract_text_with_ocr(file_path, page_num)
                        if ocr_text:
                            text_content.append(ocr_text)
                            print(f"Page {page_num + 1}: OCR fallback extracted {len(ocr_text)} characters")
                        else:
                            print(f"Page {page_num + 1}: OCR also failed")
            
            if not text_content:
                print("No text extracted from PDF")
                return None
            
            result = '\n'.join(text_content)
            print(f"Total extracted: {len(result)} characters from {len(pdf_reader.pages)} pages")
            return result
            
        except ImportError:
            print("PyPDF2 not available")
            return None
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None
    
    def _extract_text_from_html(self, file_path: str) -> Optional[str]:
        """Extract text from HTML file"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            print(f"Successfully extracted {len(text)} characters from HTML")
            return text
            
        except ImportError:
            print("BeautifulSoup not available")
            return None
        except Exception as e:
            print(f"Error extracting text from HTML: {e}")
            return None
    
    def _fallback_docx_conversion(self, file_path: str) -> Optional[str]:
        """Fallback DOCX conversion using python-docx and basic PDF creation"""
        try:
            from docx import Document
            from PIL import Image, ImageDraw, ImageFont
            
            # Read DOCX content
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Create a simple PDF using PIL
            try:
                from PIL import Image, ImageDraw, ImageFont
                
                # Create image with text
                img = Image.new('RGB', (800, 1200), color='white')
                draw = ImageDraw.Draw(img)
                
                # Use default font
                try:
                    font = ImageFont.truetype("arial.ttf", 12)
                except:
                    font = ImageFont.load_default()
                
                y_position = 50
                for line in text_content:
                    draw.text((50, y_position), line, fill='black', font=font)
                    y_position += 20
                
                # Save as PDF
                output_pdf = file_path.rsplit('.', 1)[0] + '.pdf'
                img.save(output_pdf, 'PDF')
                
                print(f"Fallback conversion successful for {file_path}")
                return output_pdf
                
            except Exception as e:
                print(f"Fallback conversion failed: {e}")
                return None
                
        except ImportError:
            print("python-docx not available for fallback conversion")
            return None
        except Exception as e:
            print(f"Fallback conversion failed: {e}")
            return None
    
    def _create_simple_pdf_from_text(self, lines: list) -> Optional[str]:
        """Create a simple PDF from text lines using PIL"""
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # Create image with text
            img = Image.new('RGB', (800, 1200), color='white')
            draw = ImageDraw.Draw(img)
            
            # Try to use a font
            try:
                font = ImageFont.truetype("arial.ttf", 12)
            except:
                font = ImageFont.load_default()
            
            y_position = 50
            for line in lines:
                if y_position < 1150:  # Keep within image bounds
                    draw.text((50, y_position), line, fill='black', font=font)
                    y_position += 20
            
            # Save as PDF
            output_pdf = f"simple_resume_{hash(str(lines))}.pdf"
            img.save(output_pdf, 'PDF')
            
            print(f"Created simple PDF from text content")
            return output_pdf
            
        except ImportError:
            print("PIL not available for simple PDF creation")
            return None
        except Exception as e:
            print(f"Error creating simple PDF: {e}")
            return None
    
    def download_and_extract_text(self, url: str) -> Optional[str]:
        """
        Download file from URL and extract text directly
        
        Args:
            url: URL to download the file from
            
        Returns:
            Extracted text content, or None if extraction fails
        """
        try:
            # Download the file
            response = requests.get(url, timeout=30)
            if response.status_code != 200:
                print(f"Failed to download file: {response.status_code}")
                return None
            
            # Determine file extension from URL or content-type
            content_type = response.headers.get('content-type', '')
            file_ext = self._get_extension_from_url(url, content_type)
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                temp_file.write(response.content)
                temp_path = temp_file.name
            
            # Detect actual file type and rename if necessary
            actual_type = self._detect_file_type(temp_path)
            if actual_type != 'unknown' and actual_type != Path(file_ext).suffix[1:]:
                # Rename file to match actual type
                new_temp_path = temp_path.rsplit('.', 1)[0] + '.' + actual_type
                try:
                    os.rename(temp_path, new_temp_path)
                    temp_path = new_temp_path
                    print(f"Renamed file to match actual type: {actual_type}")
                except Exception as e:
                    print(f"Could not rename file: {e}")
            
            # Extract text directly
            text_content = self.extract_text_from_file(temp_path)
            
            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except:
                pass
            
            return text_content
            
        except Exception as e:
            print(f"Error downloading and extracting text: {e}")
            return None
    
    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type by examining file content"""
        try:
            with open(file_path, 'rb') as f:
                # Read first 8 bytes to detect file signatures
                header = f.read(8)
                print(f"File header: {header.hex()}")
                
                # PDF signature: %PDF
                if header.startswith(b'%PDF'):
                    print("Detected as PDF")
                    return 'pdf'
                
                # DOCX signature: PK (ZIP format)
                if header.startswith(b'PK'):
                    print("Detected as DOCX")
                    return 'docx'
                
                # DOC signature: D0CF11E0 (Microsoft Compound File)
                if header.startswith(b'\xD0\xCF\x11\xE0'):
                    print("Detected as DOC")
                    return 'doc'
                
                # HTML signature: <!DOCTYPE or <html
                if header.startswith(b'<!DOCTYPE') or header.startswith(b'<html'):
                    print("Detected as HTML")
                    return 'html'
                
                # Image signatures
                if header.startswith(b'\xFF\xD8\xFF'):  # JPEG
                    print("Detected as JPEG")
                    return 'jpg'
                if header.startswith(b'\x89PNG\r\n\x1A\n'):  # PNG
                    print("Detected as PNG")
                    return 'png'
                
                # For unknown files, try to determine if it's text
                f.seek(0)
                try:
                    content = f.read(1024)
                    # Check if it's mostly printable ASCII
                    printable_count = sum(1 for byte in content if 32 <= byte <= 126 or byte in [9, 10, 13])
                    printable_ratio = printable_count / len(content)
                    print(f"Printable ratio: {printable_ratio:.2f}")
                    if printable_ratio > 0.8:  # 80% printable
                        print("Detected as TXT")
                        return 'txt'
                except:
                    pass
                
                print("Detected as unknown")
                return 'unknown'
                
        except Exception as e:
            print(f"Error detecting file type: {e}")
            return 'unknown'
    
    def _extract_text_with_ocr(self, file_path: str, page_num: int = None) -> Optional[str]:
        """Extract text from PDF using alternative methods for image-based PDFs"""
        try:
            # Try using PyMuPDF (fitz) which has better OCR capabilities
            try:
                import fitz  # PyMuPDF
                
                doc = fitz.open(file_path)
                text = ""
                
                if page_num is not None:
                    # Extract from specific page
                    page = doc[page_num]
                    text = page.get_text()
                else:
                    # Extract from all pages
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        page_text = page.get_text()
                        if page_text.strip():
                            text += page_text + "\n"
                
                doc.close()
                
                if text.strip():
                    print(f"PyMuPDF extracted {len(text)} characters")
                    return text.strip()
                else:
                    print("PyMuPDF extracted no text")
                    
            except ImportError:
                print("PyMuPDF not available")
            except Exception as e:
                print(f"PyMuPDF extraction failed: {e}")
            
            # Fallback: Try to extract any text-like content from the PDF
            try:
                import PyPDF2
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    text = ""
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            # Try to get any text-like content
                            page_text = page.extract_text()
                            if page_text:
                                # Clean up the text
                                cleaned_text = ' '.join(page_text.split())
                                if len(cleaned_text) > 10:  # Only if we got meaningful text
                                    text += cleaned_text + "\n"
                        except:
                            continue
                    
                    if text.strip():
                        print(f"Fallback extraction got {len(text)} characters")
                        return text.strip()
                        
            except Exception as e:
                print(f"Fallback extraction failed: {e}")
            
            return None
            
        except Exception as e:
            print(f"Error in OCR extraction: {e}")
            return None
    
    def _get_extension_from_url(self, url: str, content_type: str) -> str:
        """Determine file extension from URL or content-type"""
        # Try to get extension from URL
        url_lower = url.lower()
        for ext in ['.pdf', '.doc', '.docx', '.txt', '.jpg', '.jpeg', '.png']:
            if ext in url_lower:
                return ext
        
        # Try to get extension from content-type
        if 'pdf' in content_type:
            return '.pdf'
        elif 'word' in content_type or 'docx' in content_type:
            return '.docx'
        elif 'text' in content_type:
            return '.txt'
        elif 'image' in content_type:
            if 'jpeg' in content_type or 'jpg' in content_type:
                return '.jpg'
            elif 'png' in content_type:
                return '.png'
            else:
                return '.jpg'  # Default to jpg for images
        
        # Default to .txt if we can't determine
        return '.txt' 